# Project Name: ThemisHIM™ Audits
#
# Version: 1.2.7
#
# Description: ThemisHIMAudits is a Python-based tool designed for automated auditing and analysis of legal and administrative documents, particularly focusing on public procurement processes. This intelligent system processes multiple file formats, extracts relevant information, and generates structured data for easy analysis.
#
# License: CC0 1.0 Universal, CC-BY-4.0 Licence, and Apache-2.0.
#
# Author: David C Cavalcante
# Email: davcavalcante@proton.me
# LinkedIn: https://www.linkedin.com/in/hellodav/
# X: https://twitter.com/takk8is/
# Medium: https://takk8is.medium.com/
# Website: https://takk.ag/
#
# Designed by Takk™ Innovate Studio
# Takk™ Innovate Studio is at the forefront of the digital revolution, leading with a 100% Artificial Intelligence team.
#
# Research by TeleologyHI™
# Semiotics Research and Development of Hybrid Intelligence that adapts and evolves with humanity.
#
# Donations: If this project has been helpful to you, please consider making a donation to support our ongoing development of innovative tools. USDT (TRC-20) `TGpiWetnYK2VQpxNGPR27D9vfM6Mei5vNA`

import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import requests
from io import BytesIO
import pandas as pd
import pdfplumber
from docx import Document
import markdown
import pytesseract
from pdf2image import convert_from_path
import threading
import time
import logging
from pathlib import Path
import re
import csv
import tempfile
import ttkbootstrap as ttk
from typing import List, Dict, Any, Callable, Optional

log_file = os.path.join(tempfile.gettempdir(), "ThemisAudits.log")
logging.basicConfig(
    filename=log_file,
    level=logging.DEBUG,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
    bundle_dir = sys._MEIPASS
else:
    bundle_dir = os.path.dirname(os.path.abspath(__file__))


class ThemisHIMAudits:
    def __init__(self, master: ttk.Window):
        self.master = master
        master.title("ThemisHIM™ Audits")
        master.geometry("414x736")

        self.style = ttk.Style(theme="flatly")
        self._configure_styles()

        self.file_path: str = ""
        self.processed_data: List[List[str]] = []
        self.processing_thread: Optional[threading.Thread] = None

        self.create_widgets()

    def _configure_styles(self):
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure(
            "TLabel", background="#f0f0f0", foreground="#4a4a4a", font=("Inter", 12)
        )
        self.style.configure("TButton", font=("Inter", 14), padding=12, cursor="hand2")
        self.style.map(
            "TButton",
            relief=[("pressed", "flat"), ("!pressed", "flat")],
            background=[("active", "#e0e0e0"), ("!active", "#d9d9d9")],
            foreground=[("active", "#333333"), ("!active", "#4a4a4a")],
        )
        self.style.configure(
            "Accent.TButton", background="#bf7245", foreground="#ffffff"
        )
        self.style.map(
            "Accent.TButton",
            background=[("active", "#cc7949"), ("disabled", "#d9d9d9")],
            foreground=[("disabled", "#a0a0a0")],
        )
        self.style.configure(
            "Enabled.TButton", background="#bf7245", foreground="#ffffff"
        )
        self.style.map(
            "Enabled.TButton",
            background=[("active", "#cc7949"), ("!active", "#bf7245")],
            foreground=[("active", "#ffffff"), ("!active", "#ffffff")],
        )
        self.style.configure(
            "TProgressbar", background="#bf7245", troughcolor="#f0f0f0"
        )

    def create_widgets(self):
        self.main_frame = ttk.Frame(self.master, padding="30")
        self.main_frame.pack(expand=True, fill="both")

        self._create_header()
        self._create_buttons()
        self._create_progress_bar()
        self._create_footer()

    def _create_header(self):
        self.load_image()
        self.version_label = ttk.Label(
            self.main_frame,
            text="Version 1.2.7",
            font=("Inter", 14),
            foreground="#bf7245",
        )
        self.version_label.pack(pady=(10, 0))
        self.title_label = ttk.Label(
            self.main_frame, text="ThemisHIM™ Audits", font=("Inter", 32, "bold")
        )
        self.title_label.pack(pady=10)
        self.status_label = ttk.Label(
            self.main_frame, text="Ready for audit", font=("Inter", 16)
        )
        self.status_label.pack(pady=10)

    def _create_buttons(self):
        self.button_frame = ttk.Frame(self.main_frame)
        self.button_frame.pack(pady=10)

        self.audit_button = ttk.Button(
            self.button_frame,
            text="Perform Audit",
            command=self.select_file,
            style="Accent.TButton",
            width=25,
        )
        self.audit_button.pack(fill="x", pady=8)

        self.report_button = ttk.Button(
            self.button_frame,
            text="Download Report",
            command=self.generate_csv,
            state="disabled",
            style="TButton",
            width=25,
        )
        self.report_button.pack(fill="x", pady=8)

    def _create_progress_bar(self):
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.main_frame,
            variable=self.progress_var,
            maximum=100,
            length=300,
            style="TProgressbar",
        )
        self.progress_bar.pack(fill="x", pady=10)
        self.progress_bar.pack_forget()

    def _create_footer(self):
        self.signature_label = ttk.Label(
            self.main_frame,
            text="Developed by",
            font=("Inter", 12),
            foreground="#bf7245",
        )
        self.signature_label.pack(pady=(30, 0))
        self.signature_link = ttk.Label(
            self.main_frame,
            text="Takk™ Innovate Studio",
            font=("Inter", 12, "underline"),
            foreground="#bf7245",
            cursor="hand2",
        )
        self.signature_link.pack()
        self.signature_link.bind(
            "<Button-1>",
            lambda e: self.open_url("https://www.linkedin.com/in/hellodav/"),
        )

    def load_image(self):
        try:
            image_path = os.path.join(bundle_dir, "images", "icon512.png")
            if os.path.exists(image_path):
                img = Image.open(image_path).convert("RGBA")
            else:
                image_url = "https://raw.githubusercontent.com/Takk8IS/ThemisHIMAudits/main/images/icon512.png"
                response = requests.get(image_url)
                img = Image.open(BytesIO(response.content)).convert("RGBA")

            img = img.resize((250, 250), Image.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            self.image_label = ttk.Label(
                self.main_frame, image=photo, background="#f0f0f0"
            )
            self.image_label.image = photo
            self.image_label.pack(pady=10)
        except Exception as e:
            logging.error(f"Error loading image: {e}")
            self.image_label = ttk.Label(
                self.main_frame,
                text="ThemisHIM™",
                font=("Inter", 24, "bold"),
                background="#f0f0f0",
            )
            self.image_label.pack(pady=10)

    def select_file(self):
        filetypes = [
            (
                "All supported files",
                "*.pdf *.md *.rtf *.doc *.docx *.png *.jpg *.jpeg *.tiff",
            ),
            ("PDF files", "*.pdf"),
            ("Markdown files", "*.md"),
            ("RTF files", "*.rtf"),
            ("Word files", "*.doc *.docx"),
            ("Image files", "*.png *.jpg *.jpeg *.tiff"),
        ]
        file_path = filedialog.askopenfilename(filetypes=filetypes)
        if file_path:
            self.file_path = file_path
            self.status_label.config(text="Processing...")
            self.progress_var.set(0)
            self.progress_bar.pack(fill="x", pady=10)
            self.audit_button.config(state="disabled")
            self.report_button.config(state="disabled")

            self.processing_thread = threading.Thread(
                target=self.process_file, daemon=True
            )
            self.processing_thread.start()
            self.master.after(100, self.check_processing)

    def check_processing(self):
        if self.processing_thread and self.processing_thread.is_alive():
            self.master.after(100, self.check_processing)
        else:
            self.finalize_processing()

    def process_file(self):
        try:
            extension = Path(self.file_path).suffix.lower()
            processor = self.get_processor(extension)
            self.processed_data = processor(self.file_path)
        except Exception as e:
            logging.error(f"Error processing file: {e}")
            self.processed_data = []

    def finalize_processing(self):
        if self.processed_data:
            self.status_label.config(text="Audit completed")
            self.report_button.config(state="normal", style="Enabled.TButton")
        else:
            self.status_label.config(text="No data found or error occurred")

        self.progress_var.set(100)
        self.audit_button.config(state="normal")
        self.master.after(3000, self.hide_progress_bar)

    def get_processor(self, extension: str) -> Callable[[str], List[List[str]]]:
        processors = {
            ".pdf": self.process_pdf,
            ".doc": self.process_docx,
            ".docx": self.process_docx,
            ".rtf": self.process_rtf,
            ".md": self.process_md,
            ".png": self.process_image,
            ".jpg": self.process_image,
            ".jpeg": self.process_image,
            ".tiff": self.process_image,
        }
        return processors.get(extension, lambda _: [])

    def process_pdf(self, file_path: str) -> List[List[str]]:
        all_tables = []
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table in tables:
                        cleaned_table = self.clean_table(table)
                        if cleaned_table:
                            all_tables.extend(cleaned_table)
                    self.update_progress((i + 1) / total_pages * 100)
        except Exception as e:
            logging.error(f"Error processing PDF: {e}")
        return all_tables

    def process_docx(self, file_path: str) -> List[List[str]]:
        all_tables = []
        try:
            doc = Document(file_path)
            total_tables = len(doc.tables)
            for i, table in enumerate(doc.tables):
                data = [[cell.text for cell in row.cells] for row in table.rows]
                cleaned_table = self.clean_table(data)
                if cleaned_table:
                    all_tables.extend(cleaned_table)
                self.update_progress((i + 1) / total_tables * 100)
        except Exception as e:
            logging.error(f"Error processing DOCX: {e}")
        return all_tables

    def process_rtf(self, file_path: str) -> List[List[str]]:
        all_tables = []
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()
            tables = self.extract_tables_from_text(content)
            for table in tables:
                cleaned_table = self.clean_table(table)
                if cleaned_table:
                    all_tables.extend(cleaned_table)
        except Exception as e:
            logging.error(f"Error processing RTF: {e}")
        return all_tables

    def process_md(self, file_path: str) -> List[List[str]]:
        all_tables = []
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
            html = markdown.markdown(content)
            tables = self.extract_tables_from_html(html)
            for table in tables:
                cleaned_table = self.clean_table(table)
                if cleaned_table:
                    all_tables.extend(cleaned_table)
        except Exception as e:
            logging.error(f"Error processing Markdown: {e}")
        return all_tables

    def process_image(self, file_path: str) -> List[List[str]]:
        all_tables = []
        try:
            text = pytesseract.image_to_string(Image.open(file_path))
            tables = self.extract_tables_from_text(text)
            for table in tables:
                cleaned_table = self.clean_table(table)
                if cleaned_table:
                    all_tables.extend(cleaned_table)
        except Exception as e:
            logging.error(f"Error processing Image: {e}")
        return all_tables

    def extract_tables_from_text(self, text: str) -> List[List[str]]:
        lines = text.split("\n")
        tables = []
        current_table = []
        for i, line in enumerate(lines):
            if "|" in line:
                current_table.append(line.split("|"))
            elif current_table:
                tables.append(current_table)
                current_table = []
            self.update_progress((i + 1) / len(lines) * 100)
        if current_table:
            tables.append(current_table)
        return tables

    def extract_tables_from_html(self, html: str) -> List[List[str]]:
        tables = []
        pattern = r"<table.*?>(.*?)</table>"
        table_matches = re.findall(pattern, html, re.DOTALL)
        for i, table_html in enumerate(table_matches):
            rows = re.findall(r"<tr.*?>(.*?)</tr>", table_html, re.DOTALL)
            table = []
            for row in rows:
                cells = re.findall(r"<t[dh].*?>(.*?)</t[dh]>", row, re.DOTALL)
                table.append([cell.strip() for cell in cells])
            tables.append(table)
            self.update_progress((i + 1) / len(table_matches) * 100)
        return tables

    def clean_table(self, table: List[List[Any]]) -> List[List[str]]:
        cleaned_table = []
        for row in table:
            cleaned_row = [self.clean_cell(cell) for cell in row]
            if any(cleaned_row):
                cleaned_table.append(cleaned_row)
        return cleaned_table

    @staticmethod
    def clean_cell(cell: Any) -> str:
        if cell is None:
            return ""
        return re.sub(r"\s+", " ", str(cell)).strip()

    def generate_csv(self):
        if self.processed_data:
            output_file = filedialog.asksaveasfilename(
                defaultextension=".csv", filetypes=[("CSV files", "*.csv")]
            )
            if output_file:
                self.report_button.config(state="disabled")
                threading.Thread(
                    target=self.save_to_csv,
                    args=(self.processed_data, output_file),
                    daemon=True,
                ).start()

    def save_to_csv(self, data: List[List[str]], output_file: str):
        try:
            with open(output_file, "w", newline="", encoding="utf-8-sig") as file:
                writer = csv.writer(file)
                for row in data:
                    writer.writerow(row)
            self.master.after(
                0,
                lambda: self.status_label.config(
                    text=f"Report saved: {Path(output_file).name}"
                ),
            )
        except Exception as e:
            logging.error(f"Error saving CSV: {e}")
            self.master.after(
                0, lambda: self.status_label.config(text="Error saving report")
            )
        finally:
            self.master.after(
                0,
                lambda: self.report_button.config(
                    state="normal", style="Enabled.TButton"
                ),
            )

    def update_progress(self, value: float):
        self.master.after(0, lambda: self.progress_var.set(value))

    def hide_progress_bar(self):
        self.progress_bar.pack_forget()
        self.progress_var.set(0)

    @staticmethod
    def open_url(url: str):
        import webbrowser

        webbrowser.open_new(url)


def show_error(error_message):
    root = tk.Tk()
    root.withdraw()
    messagebox.showerror("Error", error_message)
    root.destroy()


def main():
    try:
        root = ttk.Window(themename="flatly")
        time.sleep(1)
        app = ThemisHIMAudits(root)
        root.mainloop()
    except Exception as e:
        error_message = f"An error occurred: {str(e)}\n\nPlease contact support."
        logging.exception("Unhandled exception:")
        show_error(error_message)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logging.exception("Fatal error in main loop")
        show_error(f"A fatal error occurred: {str(e)}\nThe application will now close.")
    finally:
        logging.info("Application closing")
